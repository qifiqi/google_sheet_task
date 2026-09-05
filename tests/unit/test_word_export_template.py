"""通用 Word JSON 导出器测试。"""

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pytest

from app.services.word_export_template import generate_word_document


def test_generate_word_document_applies_table_alignment_and_font(tmp_path):
    """通用表格统一使用微软雅黑并水平、垂直居中。"""
    output_path = tmp_path / "generic.docx"

    generate_word_document({
        "title": "通用测试文档",
        "blocks": [{
            "type": "table",
            "title": "数据表",
            "columns": ["名称", "数值"],
            "rows": [["样本", "100"]],
        }],
    }, output_path)

    table = Document(output_path).tables[0]
    assert table.alignment == WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            assert cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
            assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
            assert cell.paragraphs[0].paragraph_format.space_before.pt == 0
            assert cell.paragraphs[0].paragraph_format.space_after.pt == 0
            assert cell.paragraphs[0].runs[0].font.name == "Microsoft YaHei"


def test_generate_word_document_styles_metadata_without_changing_its_line_layout(tmp_path):
    """报告信息仍逐行输出，但标签和值具有明确的样式层级。"""
    output_path = tmp_path / "metadata.docx"

    generate_word_document({
        "title": "通用测试文档",
        "blocks": [{"type": "metadata", "items": [{"label": "报告编号", "value": "RPT-S-20260901"}]}],
    }, output_path)

    paragraph = next(item for item in Document(output_path).paragraphs if item.text.startswith("报告编号："))
    assert len(paragraph.runs) == 2
    assert paragraph.runs[0].bold is True
    assert paragraph.runs[0].font.name == "Microsoft YaHei"
    assert paragraph.runs[1].font.name == "Microsoft YaHei"


def test_generate_word_document_rejects_legacy_report_json(tmp_path):
    """旧回测专用 JSON 不再被通用渲染器接受。"""
    with pytest.raises(ValueError, match="title"):
        generate_word_document({"report_type": "RPT-S", "sections": []}, tmp_path / "legacy.docx")

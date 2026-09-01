"""通用 JSON Word 导出器。

调用方只需按 ``title + blocks`` 协议传入数据，本模块负责统一的 DOCX 样式，
不包含任何业务报告、指标或回测字段。
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


logger = logging.getLogger(__name__)


WORD_STYLE = {
    "font_name": "Microsoft YaHei",
    "page_width_inches": 8.5,
    "page_height_inches": 11,
    "margins_inches": {"left": 1.25, "right": 1.25, "top": 1, "bottom": 1},
    "normal_font_size": 10.5,
    "normal_space_after_pt": 6,
    "normal_line_spacing": 1.25,
    "headings": {
        "Title": {"font_size": 22, "color": "1F4E79", "space_before_pt": 0, "space_after_pt": 12},
        "Heading 1": {"font_size": 15, "color": "1F4E79", "space_before_pt": 14, "space_after_pt": 6},
        "Heading 2": {"font_size": 12, "color": "2F5597", "space_before_pt": 10, "space_after_pt": 6},
        "Heading 3": {"font_size": 11, "color": "2F5597", "space_before_pt": 8, "space_after_pt": 4},
    },
    "footer": {"font_size": 8, "color": "808080"},
    "table": {
        "style": "Table Grid",
        "font_size": 9,
        "header_fill": "1F4E79",
        "header_font_color": "FFFFFF",
        "alternate_fill": "F3F6FA",
        "cell_margin_twips": {"top": 90, "bottom": 90, "left": 120, "right": 120},
    },
    "metadata_label_color": "1F4E79",
    "image_width_inches": 6,
    "caption_font_size": 9,
}


def generate_word_document(document_data: dict[str, Any], output_path: str | Path) -> Path:
    """按通用 JSON 协议生成统一样式的 DOCX 文件。

    协议示例：
    ``{"title": "报告", "blocks": [{"type": "paragraph", "text": "正文"}]}``。
    """
    _validate_document_data(document_data)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _configure_document(document, str(document_data["title"]), str(document_data.get("footer") or document_data["title"]))
    _add_title(document, str(document_data["title"]))
    _render_blocks(document, document_data["blocks"])
    document.save(output_path)
    logger.info("已生成 Word 文档: %s", output_path)
    return output_path


def _validate_document_data(document_data: dict[str, Any]) -> None:
    """校验通用 Word JSON 的标题、区块及表格结构。"""
    if not isinstance(document_data, dict):
        raise ValueError("Word 文档 JSON 必须是对象")
    if not isinstance(document_data.get("title"), str) or not document_data["title"].strip():
        raise ValueError("Word 文档 JSON 必须包含非空 title")
    blocks = document_data.get("blocks")
    if not isinstance(blocks, list):
        raise ValueError("Word 文档 JSON 的 blocks 必须是数组")

    for index, block in enumerate(blocks, start=1):
        if not isinstance(block, dict):
            raise ValueError(f"第 {index} 个 Word 区块必须是对象")
        block_type = block.get("type")
        if block_type == "metadata":
            _validate_metadata_block(block, index)
        elif block_type == "heading":
            if not isinstance(block.get("text"), str) or not block["text"].strip():
                raise ValueError(f"第 {index} 个 heading 区块必须包含非空 text")
            if block.get("level", 1) not in {1, 2, 3}:
                raise ValueError(f"第 {index} 个 heading 区块的 level 仅支持 1、2、3")
        elif block_type == "paragraph":
            if not isinstance(block.get("text"), str):
                raise ValueError(f"第 {index} 个 paragraph 区块必须包含 text")
        elif block_type == "bullet_list":
            if not isinstance(block.get("items"), list) or not all(isinstance(item, str) for item in block["items"]):
                raise ValueError(f"第 {index} 个 bullet_list 区块必须包含字符串 items")
        elif block_type == "table":
            _validate_table_block(block, index)
        elif block_type == "image":
            if not isinstance(block.get("path"), str) or not block["path"].strip():
                raise ValueError(f"第 {index} 个 image 区块必须包含 path")
            if block.get("title") is not None and not isinstance(block["title"], str):
                raise ValueError(f"第 {index} 个 image 区块的 title 必须是字符串")
            if block.get("caption") is not None and not isinstance(block["caption"], str):
                raise ValueError(f"第 {index} 个 image 区块的 caption 必须是字符串")
        else:
            raise ValueError(f"第 {index} 个 Word 区块 type 不支持: {block_type}")


def _validate_metadata_block(block: dict[str, Any], index: int) -> None:
    """校验 metadata 区块中的标签和值。"""
    items = block.get("items")
    if not isinstance(items, list):
        raise ValueError(f"第 {index} 个 metadata 区块必须包含 items 数组")
    for item in items:
        if not isinstance(item, dict) or not isinstance(item.get("label"), str) or "value" not in item:
            raise ValueError(f"第 {index} 个 metadata 区块的每项必须包含 label 和 value")


def _validate_table_block(block: dict[str, Any], index: int) -> None:
    """校验通用表格的列和行数一致性。"""
    columns = block.get("columns")
    rows = block.get("rows")
    if not isinstance(columns, list) or not columns:
        raise ValueError(f"第 {index} 个 table 区块必须包含非空 columns")
    if not isinstance(rows, list) or any(not isinstance(row, list) or len(row) != len(columns) for row in rows):
        raise ValueError(f"第 {index} 个 table 区块 rows 必须与 columns 列数一致")
    if block.get("title") is not None and not isinstance(block["title"], str):
        raise ValueError(f"第 {index} 个 table 区块的 title 必须是字符串")


def _configure_document(document: Document, title: str, footer_text: str) -> None:
    """配置通用 Word 文档的页面、字体、标题和页脚样式。"""
    section = document.sections[0]
    section.page_width = Inches(WORD_STYLE["page_width_inches"])
    section.page_height = Inches(WORD_STYLE["page_height_inches"])
    margins = WORD_STYLE["margins_inches"]
    section.left_margin = Inches(margins["left"])
    section.right_margin = Inches(margins["right"])
    section.top_margin = Inches(margins["top"])
    section.bottom_margin = Inches(margins["bottom"])

    normal = document.styles["Normal"]
    _set_font_family(normal, WORD_STYLE["font_name"])
    normal.font.size = Pt(WORD_STYLE["normal_font_size"])
    normal.paragraph_format.space_after = Pt(WORD_STYLE["normal_space_after_pt"])
    normal.paragraph_format.line_spacing = WORD_STYLE["normal_line_spacing"]

    for style_name, style_config in WORD_STYLE["headings"].items():
        style = document.styles[style_name]
        _set_font_family(style, WORD_STYLE["font_name"])
        style.font.size = Pt(style_config["font_size"])
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(style_config["color"])
        style.paragraph_format.space_before = Pt(style_config["space_before_pt"])
        style.paragraph_format.space_after = Pt(style_config["space_after_pt"])
        style.paragraph_format.keep_with_next = style_name != "Title"

    bullet = document.styles["List Bullet"]
    _set_font_family(bullet, WORD_STYLE["font_name"])
    bullet.font.size = Pt(WORD_STYLE["normal_font_size"])

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(footer_text or title)
    _set_font_family(footer_run, WORD_STYLE["font_name"])
    footer_run.font.size = Pt(WORD_STYLE["footer"]["font_size"])
    footer_run.font.color.rgb = RGBColor.from_string(WORD_STYLE["footer"]["color"])


def _add_title(document: Document, title: str) -> None:
    """写入文档主标题。"""
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    _set_font_family(run, WORD_STYLE["font_name"])


def _render_blocks(document: Document, blocks: list[dict[str, Any]]) -> None:
    """按 JSON 声明顺序写入各类通用内容区块。"""
    for block in blocks:
        block_type = block["type"]
        if block_type == "metadata":
            for item in block["items"]:
                _add_metadata_item(document, str(item["label"]), str(item["value"]))
        elif block_type == "heading":
            document.add_heading(block["text"], level=block.get("level", 1))
        elif block_type == "paragraph":
            document.add_paragraph(block["text"])
        elif block_type == "bullet_list":
            for item in block["items"]:
                document.add_paragraph(item, style="List Bullet")
        elif block_type == "table":
            _add_table(document, block)
        elif block_type == "image":
            _add_image(document, block)


def _add_table(document: Document, block: dict[str, Any]) -> None:
    """按统一表头、斑马纹和居中规则写入表格。"""
    if block.get("title"):
        document.add_heading(block["title"], level=block.get("title_level", 2))
    columns = block["columns"]
    table = document.add_table(rows=1, cols=len(columns))
    table.style = WORD_STYLE["table"]["style"]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_row_repeat_as_header(table.rows[0])
    _set_row_cant_split(table.rows[0])

    for cell, value in zip(table.rows[0].cells, columns):
        _set_cell_text(cell, str(value), bold=True, color=WORD_STYLE["table"]["header_font_color"])
        _set_cell_shading(cell, WORD_STYLE["table"]["header_fill"])

    for row_index, values in enumerate(block["rows"]):
        cells = table.add_row().cells
        _set_row_cant_split(table.rows[-1])
        for cell, value in zip(cells, values):
            _set_cell_text(cell, str(value))
            if row_index % 2:
                _set_cell_shading(cell, WORD_STYLE["table"]["alternate_fill"])
    document.add_paragraph()


def _add_metadata_item(document: Document, label: str, value: str) -> None:
    """按原有逐行布局写入带标签样式的报告信息。"""
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}：")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(WORD_STYLE["metadata_label_color"])
    _set_font_family(label_run, WORD_STYLE["font_name"])
    value_run = paragraph.add_run(value)
    _set_font_family(value_run, WORD_STYLE["font_name"])


def _add_image(document: Document, block: dict[str, Any]) -> None:
    """写入图片、可选标题和居中题注。"""
    if block.get("title"):
        document.add_heading(block["title"], level=2)
    path = Path(block["path"])
    if not path.is_file():
        raise ValueError(f"图片文件不存在: {path}")
    document.add_picture(str(path), width=Inches(WORD_STYLE["image_width_inches"]))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if block.get("caption"):
        caption = document.add_paragraph(block["caption"])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.runs[0]
        run.italic = True
        run.font.size = Pt(WORD_STYLE["caption_font_size"])
        _set_font_family(run, WORD_STYLE["font_name"])


def _set_cell_text(cell: Any, text: str, bold: bool = False, color: str | None = None) -> None:
    """写入表格文本并统一字体与水平、垂直居中。"""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell, WORD_STYLE["table"]["cell_margin_twips"])
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text)
    run.bold = bold
    _set_font_family(run, WORD_STYLE["font_name"])
    run.font.size = Pt(WORD_STYLE["table"]["font_size"])
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_margins(cell: Any, margins: dict[str, int]) -> None:
    """为单元格设置对称内边距，避免文字视觉上偏向任一边。"""
    properties = cell._tc.get_or_add_tcPr()
    cell_margin = properties.first_child_found_in("w:tcMar")
    if cell_margin is None:
        cell_margin = OxmlElement("w:tcMar")
        properties.append(cell_margin)
    for side, value in margins.items():
        margin = cell_margin.find(qn(f"w:{side}"))
        if margin is None:
            margin = OxmlElement(f"w:{side}")
            cell_margin.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def _set_font_family(target: Any, font_name: str) -> None:
    """为 Word 样式或文本运行统一设置微软雅黑字体映射。"""
    target.font.name = font_name
    r_fonts = target._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), font_name)


def _set_cell_shading(cell: Any, fill: str) -> None:
    """设置单元格背景色。"""
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_row_repeat_as_header(row: Any) -> None:
    """标记首行为 Word 跨页表格的重复表头。"""
    properties = row._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    properties.append(repeat_header)


def _set_row_cant_split(row: Any) -> None:
    """保持单个表格行完整，避免一行内容被断在两页。"""
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))

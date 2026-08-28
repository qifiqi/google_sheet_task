"""量化策略回测报告 DOCX 模板。

本模块只负责将已计算的报告 JSON 排版成 DOCX，不计算任何回测指标。
"""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


logger = logging.getLogger(__name__)
# 报告结构、样式和校验规则均由该 JSON 配置；渲染器只负责通用解释。
DEFAULT_TEMPLATE_PATH = Path(__file__).with_name("report_templates") / "strategy_backtest_report.json"


def generate_strategy_backtest_report(
    report_data: dict[str, Any],
    output_path: str | Path,
    template_path: str | Path = DEFAULT_TEMPLATE_PATH,
) -> Path:
    """按 JSON 模板和已计算的数据生成 DOCX 报告。"""
    # 模板决定“渲染什么”，report_data 仅提供“渲染的数据”。
    template = _load_template(template_path)
    _validate_report_data(report_data, template)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _configure_document(document, template["styles"], report_data)
    _render_blocks(document, report_data, template, template["blocks"])

    document.save(output_path)
    logger.info("已生成策略回测报告: %s", output_path)
    return output_path


def _load_template(template_path: str | Path) -> dict[str, Any]:
    """处理_load_template相关逻辑。"""
    path = Path(template_path)
    if not path.is_file():
        raise ValueError(f"报告 JSON 模板不存在: {path}")
    try:
        with path.open(encoding="utf-8") as template_file:
            return json.load(template_file)
    except json.JSONDecodeError as error:
        raise ValueError(f"报告 JSON 模板格式无效: {path}") from error


def _validate_report_data(report_data: dict[str, Any], template: dict[str, Any]) -> None:
    """处理_validate_report_data相关逻辑。"""
    validation = template["validation"]
    required_keys = set(validation["required_fields"])
    missing_keys = required_keys - report_data.keys()
    if missing_keys:
        raise ValueError(f"report_data 缺少字段: {', '.join(sorted(missing_keys))}")

    if report_data["report_type"] not in set(validation["report_types"]):
        raise ValueError(f"report_type 仅支持 {'、'.join(validation['report_types'])}")

    metadata_keys = set(validation["required_metadata_fields"])
    missing_metadata = metadata_keys - report_data["metadata"].keys()
    if missing_metadata:
        raise ValueError(f"metadata 缺少字段: {', '.join(sorted(missing_metadata))}")

    # 条件校验（例如 RPT-S 单产品、RPT-M 多产品）由 JSON 配置驱动。
    for condition in validation.get("conditions", []):
        if _resolve_value(report_data, condition["field"]) != condition["equals"]:
            continue
        for field in condition.get("required_fields", []):
            if not _resolve_optional_value(report_data, field):
                raise ValueError(f"{report_data['report_type']} 报告必须传入 {field}")
        for field in condition.get("forbidden_fields", []):
            if _resolve_optional_value(report_data, field):
                raise ValueError(f"{report_data['report_type']} 报告不应传入 {field}")
        row_count_rule = condition.get("table_row_count")
        if row_count_rule:
            table_data = _resolve_value(report_data, row_count_rule["field"])
            if len(table_data.get("rows", [])) != row_count_rule["equals"]:
                raise ValueError(
                    f"{report_data['report_type']} 报告的 {row_count_rule['field']} "
                    f"必须包含 {row_count_rule['equals']} 行"
                )

    weight_allocation = report_data.get("weight_allocation")
    if weight_allocation:
        _validate_table_data(weight_allocation, "weight_allocation")

    for section in report_data["sections"]:
        if not section.get("title") or not section.get("subsections"):
            raise ValueError("每个 section 必须包含 title 和非空的 subsections")
        for subsection in section["subsections"]:
            table = subsection.get("table")
            _validate_table_data(table, subsection.get("title", section["title"]))


def _validate_table_data(table_data: dict[str, Any] | None, table_name: str) -> None:
    """处理_validate_table_data相关逻辑。"""
    if not table_data or not table_data.get("columns"):
        raise ValueError(f"{table_name} 必须包含 table.columns")
    column_count = len(table_data["columns"])
    if any(len(row) != column_count for row in table_data.get("rows", [])):
        raise ValueError(f"表格列数不一致: {table_name}")


def _render_blocks(
    document: Document,
    report_data: dict[str, Any],
    template: dict[str, Any],
    blocks: list[dict[str, Any]],
) -> None:
    """处理_render_blocks相关逻辑。"""
    for block in blocks:
        block_type = block["type"]
        # 每一种 block 都对应 JSON 中的一种声明式报告组件。
        if block_type == "conditional":
            condition = block["when"]
            if _resolve_value(report_data, condition["field"]) == condition["equals"]:
                _render_blocks(document, report_data, template, block["blocks"])
        elif block_type == "title":
            title = document.add_paragraph(style="Title")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.add_run(str(_resolve_value(report_data, block["source"])))
        elif block_type == "metadata":
            _add_metadata(document, report_data, template)
        elif block_type == "heading":
            document.add_heading(block["text"], level=block["level"])
        elif block_type == "text":
            document.add_paragraph(block["text"])
        elif block_type == "table":
            _add_table(
                document,
                _resolve_value(report_data, block["source"]),
                template["styles"]["table"],
                header_fill=block.get("header_fill", True),
            )
        elif block_type == "sections":
            for section in _resolve_value(report_data, block["source"]):
                document.add_heading(section["title"], level=1)
                for subsection in section["subsections"]:
                    if subsection.get("title"):
                        document.add_heading(subsection["title"], level=2)
                    _add_table(document, subsection["table"], template["styles"]["table"])
        elif block_type == "charts":
            for chart in _resolve_value(report_data, block["source"]):
                document.add_heading(chart["title"], level=2)
                _add_chart(document, chart, template["styles"]["chart"])
        elif block_type == "notes":
            for title, notes in _resolve_value(report_data, block["source"]):
                document.add_heading(title, level=2)
                for note in notes:
                    document.add_paragraph(note)
        elif block_type == "paragraphs":
            for paragraph in _resolve_value(report_data, block["source"]):
                document.add_paragraph(paragraph)
        else:
            raise ValueError(f"不支持的报告 JSON 区块类型: {block_type}")


def _resolve_value(data: dict[str, Any], source: str) -> Any:
    """处理_resolve_value相关逻辑。"""
    value: Any = data
    # 支持 metadata.report_id 这类点路径，避免将字段位置固化在 Python 中。
    for key in source.split("."):
        if not isinstance(value, dict) or key not in value:
            raise ValueError(f"report_data 缺少 JSON 模板需要的字段: {source}")
        value = value[key]
    return value


def _resolve_optional_value(data: dict[str, Any], source: str) -> Any:
    """处理_resolve_optional_value相关逻辑。"""
    try:
        return _resolve_value(data, source)
    except ValueError:
        return None


def _configure_document(document: Document, styles: dict[str, Any], report_data: dict[str, Any]) -> None:
    """处理_configure_document相关逻辑。"""
    document_style = styles["document"]
    section = document.sections[0]
    section.page_width = Inches(document_style["page_width_inches"])
    section.page_height = Inches(document_style["page_height_inches"])
    margins = document_style["margins_inches"]
    section.left_margin = Inches(margins["left"])
    section.right_margin = Inches(margins["right"])
    section.top_margin = Inches(margins["top"])
    section.bottom_margin = Inches(margins["bottom"])

    normal = document.styles["Normal"]
    normal.font.name = document_style["font_name"]
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), document_style["font_name"])
    normal.font.size = Pt(document_style["normal_font_size"])
    normal.paragraph_format.space_after = Pt(document_style["normal_space_after_pt"])
    normal.paragraph_format.line_spacing = document_style["normal_line_spacing"]

    # 标题样式与“标题不单独留在页尾”的规则均从 JSON 读取。
    for style_name, style_config in styles["headings"].items():
        style = document.styles[style_name]
        style.font.name = document_style["font_name"]
        style._element.rPr.rFonts.set(qn("w:eastAsia"), document_style["font_name"])
        style.font.size = Pt(style_config["font_size"])
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(style_config["color"])
        style.paragraph_format.space_before = Pt(style_config["space_before_pt"])
        style.paragraph_format.space_after = Pt(style_config["space_after_pt"])
        style.paragraph_format.keep_with_next = style_config.get("keep_with_next", False)

    footer = section.footer.paragraphs[0]
    footer_style = styles["footer"]
    footer.alignment = _alignment(footer_style["alignment"])
    footer_run = footer.add_run(footer_style["text"].format(**report_data))
    footer_run.font.name = document_style["font_name"]
    footer_run._element.rPr.rFonts.set(qn("w:eastAsia"), document_style["font_name"])
    footer_run.font.size = Pt(footer_style["font_size"])
    footer_run.font.color.rgb = RGBColor.from_string(footer_style["color"])


def _add_metadata(document: Document, report_data: dict[str, Any], template: dict[str, Any]) -> None:
    """处理_add_metadata相关逻辑。"""
    metadata_style = template["styles"]["metadata"]
    for field in template["metadata_fields"]:
        paragraph = document.add_paragraph()
        paragraph.alignment = _alignment(metadata_style["alignment"])
        paragraph.add_run(f"{field['label']}：{_resolve_value(report_data, field['source'])}")


def _add_table(
    document: Document,
    table_data: dict[str, Any],
    table_style: dict[str, Any],
    header_fill: bool = True,
) -> None:
    """处理_add_table相关逻辑。"""
    table = document.add_table(rows=1, cols=len(table_data["columns"]))
    table.style = table_style["style"]
    table.autofit = True
    # Word 需要在 OOXML 中显式标记，才能在表格跨页时重复第一行表头。
    if table_style.get("repeat_header_row", False):
        _set_row_repeat_as_header(table.rows[0])
    # 防止同一数据行内容被拆到两页；长表仍可在行与行之间分页。
    if not table_style.get("allow_row_break_across_pages", True):
        for row in table.rows:
            _set_row_cant_split(row)
    header_cells = table.rows[0].cells
    for cell, value in zip(header_cells, table_data["columns"]):
        _set_cell_text(
            cell,
            str(value),
            table_style,
            bold=True,
            color=table_style["header_font_color"] if header_fill else None,
        )
        if header_fill:
            _set_cell_shading(cell, table_style["header_fill"])

    for row_index, values in enumerate(table_data.get("rows", [])):
        cells = table.add_row().cells
        if not table_style.get("allow_row_break_across_pages", True):
            _set_row_cant_split(table.rows[-1])
        for cell, value in zip(cells, values):
            _set_cell_text(cell, str(value), table_style)
            if row_index % 2:
                _set_cell_shading(cell, table_style["alternate_fill"])

    document.add_paragraph()


def _set_cell_text(
    cell: Any,
    text: str,
    table_style: dict[str, Any],
    bold: bool = False,
    color: str | None = None,
) -> None:
    """处理_set_cell_text相关逻辑。"""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = _alignment(table_style["alignment"])
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = table_style["font_name"]
    run._element.rPr.rFonts.set(qn("w:eastAsia"), table_style["font_name"])
    run.font.size = Pt(table_style["font_size"])
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_shading(cell: Any, fill: str) -> None:
    """处理_set_cell_shading相关逻辑。"""
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_row_repeat_as_header(row: Any) -> None:
    """标记首行为 Word 跨页表格的重复表头。"""
    properties = row._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    properties.append(repeat_header)


def _set_row_cant_split(row: Any) -> None:
    """保持单个表格行完整，避免一行内容被断在两页。"""
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def _add_chart(document: Document, chart: dict[str, Any], chart_style: dict[str, Any]) -> None:
    """处理_add_chart相关逻辑。"""
    image_path = chart.get("image_path")
    if image_path:
        path = Path(image_path)
        if not path.is_file():
            raise ValueError(f"图表文件不存在: {path}")
        document.add_picture(str(path), width=Inches(chart_style["width_inches"]))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        placeholder = document.add_paragraph("【图表待传入】")
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        placeholder.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    if chart.get("caption"):
        caption = document.add_paragraph(chart["caption"])
        caption.alignment = _alignment(chart_style["caption_alignment"])
        caption.runs[0].italic = True
        caption.runs[0].font.size = Pt(chart_style["caption_font_size"])


def _alignment(value: str) -> WD_ALIGN_PARAGRAPH:
    """处理_alignment相关逻辑。"""
    alignments = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    if value not in alignments:
        raise ValueError(f"不支持的文本对齐方式: {value}")
    return alignments[value]
